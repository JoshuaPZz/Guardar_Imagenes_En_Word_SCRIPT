import os
from docx import Document
from docx.shared import Cm

# Ruta de la carpeta que contiene las imágenes
carpeta_imagenes = 'San Inductín'

# Ruta de salida del documento de Word
documento_word = 'sanInducTin.docx'

# Crear un nuevo documento de Word
doc = Document()

# Ajustar márgenes de la página para más espacio (opcional)
section = doc.sections[0]
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)

# Tamaño deseado para las imágenes
ancho = Cm(7.5)
alto = Cm(5)

# Número de imágenes por fila y filas por página
imagenes_por_fila = 2
filas_por_pagina = 4
imagenes_por_pagina = imagenes_por_fila * filas_por_pagina

# Contador para imágenes y control de filas
contador = 0
fila_actual = 0

# Recorrer todas las imágenes en la carpeta
for nombre_imagen in sorted(os.listdir(carpeta_imagenes)):  # Ordenar para consistencia
    if contador >= 296:
        break

    if nombre_imagen.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
        ruta_imagen = os.path.join(carpeta_imagenes, nombre_imagen)

        # Si es la primera imagen de la fila, se crea un nuevo párrafo
        if contador % imagenes_por_fila == 0:
            parrafo_actual = doc.add_paragraph()

        # Añadir imagen al párrafo actual
        run = parrafo_actual.add_run()
        run.add_picture(ruta_imagen, width=ancho, height=alto)
        run.add_text(" " * 5)  # Espacio entre imágenes

        contador += 1

        # Control de filas
        if contador % imagenes_por_fila == 0:
            fila_actual += 1

        # Si alcanzamos el número de imágenes por página, agregamos un salto de página
        if contador % imagenes_por_pagina == 0 and contador < 296:
            doc.add_page_break()

# Guardar el documento
doc.save(documento_word)
print(f"Se insertaron {contador} imágenes en {documento_word}!")
